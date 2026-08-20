#!/usr/bin/env python3
"""
Validate a TIG article set against docs/tyler-insurance-group-article-guide.md.

Checks, per article:
  HTML  tag balance; zero em dashes; exactly one Bluegrass backlink; no local
        references outside that link's anchor; agency voice in narrative prose
        (FAQ questions and quoted reader speech are exempt, since those are
        deliberately written in the reader's voice); word count.
  DOCX  zero Heading 1; first heading is Heading 2; two or more body H3s before
        the FAQ; no heading level jumps; no ALL-CAPS pseudo-headings; no empty
        paragraphs.
  SVG   well-formed XML; 1200x630; TIG palette only, no Bluegrass coral/green.
  PNG   present and correctly sized.

Usage:  python3 docs/check-tig.py [dir]
"""
import re, sys, glob, os, xml.dom.minidom
from html.parser import HTMLParser

D = sys.argv[1] if len(sys.argv) > 1 else "dist/tig-wave1"

TIG_OK = {"#dbcf86", "#b0a154", "#f6f4ea", "#efe8cf", "#7e8082", "#5c5e60",
          "#ffffff", "#fff", "#2a2620", "#5f594f", "#e4e0d4", "#8a8a8a"}
BANNED_HEX = {"#d05528", "#b3431d", "#3a7d52"}          # Bluegrass coral / green
LOCAL = r'\b(Lexington|Bluegrass|Fayette|Baptist Health|UK HealthCare|Saint Joseph|Winchester|Nicholasville|Jessamine|Clark County)\b'

VOID = {'area','base','br','col','embed','hr','img','input','link','meta','param','source','track','wbr'}
class Bal(HTMLParser):
    def __init__(s): super().__init__(convert_charrefs=True); s.st=[]; s.e=[]
    def handle_starttag(s,t,a):
        if t not in VOID: s.st.append(t)
    def handle_endtag(s,t):
        if t in VOID: return
        if s.st and s.st[-1]==t: s.st.pop()
        else: s.e.append(t)

def text_of(src, drop_faq=False):
    """Visible text. drop_faq removes the FAQ block and any quoted sentence,
    both of which are legitimately written in the reader's first person."""
    if drop_faq:
        src = re.sub(r'<h2>Frequently asked questions</h2>.*?(?=<div class="cta">)', '', src, flags=re.S)
        src = re.sub(r'"[^"]{10,240}"', '', src)
    src = re.sub(r'<(script|style|head)\b.*?</\1>', '', src, flags=re.S)
    return re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', src))

def main():
    fails = []
    slugs = sorted(os.path.basename(p)[:-5] for p in glob.glob(f"{D}/*.html"))
    if not slugs:
        print(f"no articles in {D}"); return 1
    print(f"{'article':40s} {'words':>5} {'H3':>3} {'links':>5}  status")
    print("-" * 78)

    for s in slugs:
        bad = []
        src = open(f"{D}/{s}.html", encoding="utf-8").read()

        b = Bal(); b.feed(src); b.close()
        if b.e or b.st: bad.append(f"tags {b.e}{[x for x in b.st]}")

        body = text_of(src)
        if body.count("—") or "&mdash;" in src: bad.append("em dash")

        links = re.findall(r'href="(https://www\.bluegrassmedicarehelp\.com[^"]*)"', src)
        if len(links) != 1: bad.append(f"{len(links)} bluegrass links, want 1")

        # local references are only allowed inside the sanctioned backlink anchor
        outside = re.sub(r'<a [^>]*bluegrassmedicarehelp[^>]*>.*?</a>', '', src, flags=re.S)
        loc = sorted(set(re.findall(LOCAL, text_of(outside))))
        if loc: bad.append(f"local refs {loc}")

        # agency voice, narrative prose only
        narr = text_of(src, drop_faq=True)
        fp = sorted(set(re.findall(r"\b(I|I'm|I've|my|mine)\b", narr)))
        if fp: bad.append(f"first person {fp}")
        if re.search(r'\bAustin\b', text_of(re.sub(r'<p class="(byline|author)">.*?</p>', '', src, flags=re.S))):
            bad.append("Austin named in body")

        # docx
        from docx import Document
        doc = Document(f"{D}/{s}.docx")
        heads = [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith("Heading")]
        lv = [int(h[0].split()[-1]) for h in heads]
        faq_i = next((i for i, h in enumerate(heads) if "Frequently asked" in h[1]), len(heads))
        h3 = sum(1 for h in heads[:faq_i] if h[0] == "Heading 3")
        if any(h[0] == "Heading 1" for h in heads): bad.append("docx has H1")
        if not heads or heads[0][0] != "Heading 2": bad.append("docx first heading not H2")
        if any(lv[i] - lv[i-1] > 1 for i in range(1, len(lv))): bad.append("docx level jump")
        if h3 < 2: bad.append(f"docx only {h3} body H3")
        if any(p.text.strip().isupper() and len(p.text.strip()) > 3 for p in doc.paragraphs):
            bad.append("ALL-CAPS pseudo-heading")
        if any(not p.text.strip() for p in doc.paragraphs): bad.append("empty paragraph")

        # svg
        sv = open(f"{D}/{s}.svg", encoding="utf-8").read()
        try: xml.dom.minidom.parseString(sv)
        except Exception as e: bad.append(f"svg malformed {e}")
        if 'width="1200"' not in sv or 'height="630"' not in sv: bad.append("svg not 1200x630")
        hexes = {h.lower() for h in re.findall(r'#[0-9a-fA-F]{3,6}', sv)}
        if hexes & BANNED_HEX: bad.append(f"bluegrass palette {hexes & BANNED_HEX}")
        stray = hexes - TIG_OK
        if stray: bad.append(f"off-palette {sorted(stray)}")

        # png
        png = f"{D}/{s}.png"
        if not os.path.exists(png): bad.append("png missing")
        else:
            from PIL import Image
            w, hh = Image.open(png).size
            if (w, hh) != (1200, 630): bad.append(f"png {w}x{hh}")

        if bad: fails.append((s, bad))
        print(f"{s:40s} {len(body.split()):5d} {h3:3d} {len(links):5d}  {'OK' if not bad else 'FAIL'}")

    if fails:
        print()
        for s, bad in fails:
            print(f"  {s}")
            for x in bad: print(f"     - {x}")
        print(f"\n{len(fails)} article(s) failing")
        return 1
    print(f"\nall {len(slugs)} articles pass")
    return 0

if __name__ == "__main__":
    sys.exit(main())
