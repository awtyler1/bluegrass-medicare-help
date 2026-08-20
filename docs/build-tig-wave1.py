#!/usr/bin/env python3
"""
Build the TIG Wave 1 article set for AEP 2026 / plan year 2027.

Per docs/tyler-insurance-group-article-guide.md, each article ships as:
  1. a self-contained styled HTML doc (paste into the TIG editor)
  2. an editable Word .docx (title line + body, no H1 in the body)
  3. a 1200x630 illustration in the TIG palette, as SVG and PNG

One content source (BLOCKS) feeds all three so they cannot drift.
Every figure here is corroboration-verified; see docs/tig-editorial-plan-aep-2027.md section 9.

Usage:  python3 docs/build-tig-wave1.py [outdir]
"""
import os, re, sys, html as _html

OUT = sys.argv[1] if len(sys.argv) > 1 else "dist/tig-wave1"

# ---------------------------------------------------------------- brand
GOLD, GOLD_D, GOLD_T = "#dbcf86", "#b0a154", "#f6f4ea"
GRAY, GRAY_D = "#7e8082", "#5c5e60"
INK, MUTE, LINE, WHITE = "#2a2620", "#5f594f", "#e4e0d4", "#ffffff"

PHONE = "(859) 618-6443"
UPDATED = "August 2026"

DISCLAIMER = (
    "This article is general information, not advice for your specific situation, and Medicare "
    "rules and figures change. Figures cited are from CMS and the Social Security Administration "
    "and were verified in August 2026; the 2027 Part B premium is a Trustees Report projection "
    "until CMS confirms it. Rules and available plans vary by state and by county. Tyler "
    "Insurance Group is not connected with or endorsed by the United States government or the "
    "federal Medicare program. We do not offer every plan available in your area. Please contact "
    "Medicare.gov, 1-800-MEDICARE, or your State Health Insurance Assistance Program (SHIP) to "
    "get information on all of your options."
)
AUTHOR_LINE = ("Written by Austin Tyler, a licensed insurance agent with Tyler Insurance Group.")


# ---------------------------------------------------------------- content
def A(**kw):
    return kw

ARTICLES = [

# ============================================================ 1.1
A(
slug="medicare-advantage-plan-ending-2027",
title="Medicare Advantage Plan Ending in 2027? Your Options",
h1="Your Medicare Advantage Plan Is Ending in 2027: What the September Letter Means and What to Do Next",
desc="About 600,000 Humana members will get a plan non-renewal letter in September. What it means, the two rights it triggers, and the one that expires in 63 days.",
art="letter",
blocks=[
 ("p","In September, roughly <strong>600,000 people</strong> will open an envelope telling them their Medicare Advantage plan will not exist next year. Humana announced on its July 29, 2026 earnings call that it is exiting plans covering about 8 percent of its 7.2 million Medicare Advantage members for 2027, and it is not the only carrier trimming its map."),
 ("p","If one of those letters is sitting on your kitchen table, here is the part almost nobody tells you: <strong>losing your plan hands you a right you will probably never get again.</strong> It is worth real money, and it expires."),
 ("qa","Your plan ending triggers two things. A <strong>Special Enrollment Period</strong> to pick new coverage, and a <strong>guaranteed-issue right</strong> to buy certain Medicare Supplement (Medigap) policies with <strong>no health questions asked</strong>. That second one has a <strong>63-day</strong> clock on it. Doing nothing is the one choice that can cost you both."),

 ("h2","First, read what the letter actually says"),
 ("p","Three different letters get mailed in the fall and people mix them up constantly. The wording decides which rights you have, so find your letter and check which one you are holding."),
 ("h3","A plan non-renewal or termination"),
 ("p","Your specific plan is being discontinued in your county for next year. This is the one that triggers everything below. The letter will use language like \"will no longer be offered\" or \"is being discontinued.\""),
 ("h3","The carrier is leaving your county entirely"),
 ("p","Same practical effect for you, and the same rights. Worth knowing the difference only because a carrier can pull one plan while keeping others in your area, which changes what is available to switch into locally."),
 ("h3","An Annual Notice of Change"),
 ("p","This is not a termination. Your plan continues, but the premium, drug list, provider network or copays are changing. Everyone in a Medicare Advantage or Part D plan gets one by September 30. It deserves a careful read, but it does not give you a guaranteed-issue Medigap right."),

 ("h2","The right that expires, and the two lists"),
 ("p","When a Medicare Advantage plan terminates, federal law gives you a guaranteed-issue right to buy a Medigap policy. An insurer must sell it to you, cannot ask about your health, and cannot charge you more because of your medical history. Outside of a protected window like this one, most states let insurers underwrite Medigap applications and turn people down."),
 ("callout","The 63-day clock.","You generally have 63 days from when your coverage ends to apply under this protection. Keep the termination letter. Insurers routinely ask for proof of the qualifying event, and that letter is the proof."),
 ("p","Which Medigap plans an insurer must sell you depends on <strong>when you first became eligible for Medicare</strong>, because of a 2015 law called MACRA. There are two answers and most articles only print the old one:"),
 ("table","Guaranteed-issue Medigap plans on plan termination",
   ["If you first became eligible for Medicare","Plans an insurer must sell you"],
   [["Before January 1, 2020","A, B, C, F, K, L"],
    ["On or after January 1, 2020","A, B, D, G, K, L"]]),
 ("p","If you are turning 65 now, you are always in the second row. Plans C and F were closed to newly eligible people in 2020 because they covered the Part B deductible, and <strong>Plans D and G took their place</strong> as the guaranteed-issue options. Anyone who tells you a terminated member can get Plan F is working from an outdated list."),
 ("p","These are the federal minimums. <strong>Several states are more generous</strong>, with broader guaranteed-issue windows or year-round protections, so check your own state insurance department before assuming the federal floor is all you get."),

 ("h2","What happens if you ignore the letter"),
 ("p","You will usually be automatically enrolled in another plan from the same carrier, or dropped back to Original Medicare with no drug coverage. Neither outcome is chosen for your benefit, and both quietly burn the guaranteed-issue window while you wait."),
 ("p","If you land in Original Medicare with no Part D plan, you also start accruing a <strong>late enrollment penalty</strong> that is added to your drug premium permanently once you do sign up."),

 ("h2","Your four options, side by side"),
 ("table","What you can do, and what each one costs you",
   ["Option","What to know"],
   [["Another Medicare Advantage plan","Usually the easiest switch. Check that your doctors, hospital and prescriptions are all covered under the new plan, not just the premium."],
    ["Original Medicare plus a Medigap policy","This is where the guaranteed-issue right applies. Higher monthly premium, no network, very predictable costs. You will also want a separate Part D plan."],
    ["Original Medicare plus Part D only","Lowest premium, but Original Medicare has no annual out-of-pocket maximum. A serious year can get expensive."],
    ["Do nothing","You get whatever you are assigned, and the guaranteed-issue window closes."]]),

 ("h2","What to do this month"),
 ("ol",["Find the letter and confirm it says your plan is being discontinued, not merely changing.",
        "Write down the date your coverage ends. That starts the 63-day clock.",
        "Decide whether you want a network or no network. That single question separates Medicare Advantage from Medigap and it is easier to answer than comparing dozens of plans.",
        "If Medigap is even a maybe, price it now while no one can ask about your health. You can always decline it. You cannot always get it back.",
        "Check that any Advantage plan you are considering covers your doctors, your hospital and every prescription you take, by name."]),
 ("p","Our licensed agents run that comparison for people in this exact situation every fall, at no cost. For a look at how this plays out in one specific market, our local Kentucky team walked through it in their guide to <a href=\"https://www.bluegrassmedicarehelp.com/articles/medicare-advantage-plan-ending-kentucky/\">what to do when a Medicare Advantage plan ends</a>."),

 ("recap",["About 600,000 Humana members are affected for 2027, with letters arriving in September.",
           "A plan termination gives you a Special Enrollment Period and a guaranteed-issue right to buy Medigap with no health questions.",
           "You generally have 63 days from the end of coverage. Keep the letter as proof.",
           "Eligible before January 1, 2020: Plans A, B, C, F, K, L. Eligible on or after: Plans A, B, D, G, K, L.",
           "Doing nothing usually means an automatic assignment and a closed window."]),

 ("faq",[
  ("Is my Medicare Advantage plan really ending, or just changing?",
   "Check the letter's wording. A termination or non-renewal notice says the plan will no longer be offered, and that triggers a Special Enrollment Period and guaranteed-issue Medigap rights. An Annual Notice of Change means the plan continues with different costs or benefits next year, which does not create those rights. Everyone in a Medicare Advantage or Part D plan receives an Annual Notice of Change by September 30."),
  ("Can an insurance company turn me down for Medigap if my plan was discontinued?",
   "No, not during the guaranteed-issue window. When your Medicare Advantage plan terminates, insurers must sell you certain Medigap policies regardless of your health history, cannot apply a pre-existing condition waiting period, and cannot charge you more because of your medical history. You generally have 63 days from the end of your coverage to apply, and you should keep the termination letter as proof of the qualifying event."),
  ("Which Medigap plans can I buy if my plan is ending?",
   "It depends on when you first became eligible for Medicare. If that was before January 1, 2020, the guaranteed-issue plans are A, B, C, F, K and L. If you became eligible on or after January 1, 2020, they are A, B, D, G, K and L, because a 2015 law called MACRA closed Plans C and F to newly eligible people and substituted D and G. Some states offer broader rights than this federal minimum, so check with your state insurance department."),
  ("What happens if I do nothing when my plan ends?",
   "You will typically be automatically enrolled in another plan from the same carrier, or returned to Original Medicare without drug coverage. Neither is chosen for your benefit, and both let the 63-day guaranteed-issue window close. If you end up without Part D coverage, you may also begin accruing a late enrollment penalty that is permanently added to your drug premium once you enroll."),
  ("Do I have to wait for the Annual Enrollment Period to switch?",
   "No. A plan termination gives you a Special Enrollment Period, so you are not limited to the October 15 to December 7 window. You can also use the Annual Enrollment Period if that timing suits you better, but the guaranteed-issue Medigap right runs on its own 63-day clock tied to when your coverage ends, not to the enrollment calendar."),
 ]),
]),

# ============================================================ 1.2
A(
slug="medicare-changes-2027",
title="Medicare Changes in 2027: What You'll Actually Pay",
h1="Medicare in 2027: Every Change That Affects What You Pay",
desc="The 2027 numbers in one place: Part B, the Part D deductible and out-of-pocket cap, negotiated drug prices, and the plan exits reshaping the map.",
art="calendar",
blocks=[
 ("p","Most write-ups of a new Medicare year bury the numbers under six paragraphs of preamble. Here they are first."),
 ("table","What changes on January 1, 2027",
   ["","2026","2027"],
   [["Part B standard premium","$202.90","about $209.50 (projected)"],
    ["Part D standard deductible","$615","$700"],
    ["Part D out-of-pocket cap","$2,100","$2,400"],
    ["Negotiated drug prices","first 10 drugs in effect","second round of 15 drugs takes effect"]]),
 ("p","Now the part that matters, which is what each of those does to a real budget."),

 ("h2","Part B: a small increase, and a projection"),
 ("p","The 2026 Medicare Trustees Report projects a standard Part B premium of about <strong>$209.50</strong> for 2027, up from the confirmed $202.90. That is roughly a 3 percent increase, which is modest by recent standards."),
 ("callout","Treat that number as a forecast, not a fact.","CMS confirms the real Part B premium in the fall, usually in November, which lands in the middle of the Annual Enrollment Period. Anyone publishing $209.50 as settled before then is quoting a projection. We will update this figure when it is confirmed."),
 ("p","Higher earners pay more through IRMAA, an income-related surcharge based on your tax return from two years earlier. If your income dropped since then because you retired, sold a business or lost a spouse, you can ask Social Security to use current income instead."),

 ("h2","Part D: the deductible rises faster than the cap"),
 ("h3","The $700 deductible"),
 ("p","The standard Part D deductible goes from $615 to <strong>$700</strong>, about a 14 percent jump. This is the change most people will actually feel, because everyone who fills a covered prescription runs through the deductible before the plan starts paying its share. Plans may set a lower deductible than the standard, and some set none at all, so check your own plan rather than assuming $700."),
 ("h3","The $2,400 out-of-pocket cap"),
 ("p","The annual cap on what you pay out of pocket for covered Part D drugs rises from $2,100 to <strong>$2,400</strong>. Once you hit it, you pay nothing more for covered drugs for the rest of the calendar year."),
 ("p","A rising cap sounds like bad news and mostly is not. The cap exists to stop a catastrophic year, and most people never reach it. If you take an expensive specialty drug, you will reach it and the $300 increase is real. If you take four generics, the deductible change matters to you and the cap almost certainly does not."),

 ("h2","Negotiated drug prices reach a second group of drugs"),
 ("p","Negotiated prices for a second round of 15 Part D drugs take effect January 1, 2027, at an average reduction of roughly 44 percent off list price. The group covers drugs used by about 5.3 million Part D enrollees and includes widely used products such as the semaglutide medicines, Trelegy Ellipta, Linzess, Ibrance, Xtandi and Austedo."),
 ("callout","A lower negotiated price does not automatically lower your copay.","What you pay is set by your plan's formulary tier, your deductible and which phase of the benefit you are in. Some people will see a real drop. Others will see nothing change. The only way to know is to look up your own drugs in your own plan for 2027."),

 ("h2","Fewer plans on the map"),
 ("p","Carriers are continuing to shrink their Medicare Advantage footprints. Humana said in July 2026 that it will exit plans covering roughly 600,000 members for 2027, following a similar reduction the year before, and other carriers have signaled a shift from growth to margin."),
 ("p","The practical consequence is that a meaningful number of people will get a letter this September saying their plan will not exist next year. That letter carries rights with a deadline attached, which our local Kentucky team covers in their guide to <a href=\"https://www.bluegrassmedicarehelp.com/articles/medicare-advantage-plan-ending-kentucky/\">what to do when a Medicare Advantage plan ends</a>."),

 ("h2","What to actually do about all this"),
 ("ol",["Read your Annual Notice of Change when it arrives by September 30. It tells you what your specific plan is doing, which matters more than any national average.",
        "Look up every prescription you take in your plan's 2027 formulary during the Annual Enrollment Period, October 15 to December 7.",
        "Confirm your doctors and hospital are still in network for 2027. Networks change on January 1 without asking you.",
        "If your income has dropped since the tax return Social Security is using, file to have the IRMAA surcharge reconsidered."]),

 ("recap",["Part B is projected at about $209.50, up from $202.90, with CMS confirming in the fall.",
           "The Part D deductible rises to $700 from $615, which is the change most people will feel.",
           "The Part D out-of-pocket cap rises to $2,400 from $2,100.",
           "Negotiated prices on a second group of 15 drugs start January 1, at about 44 percent off list.",
           "Plan exits continue, so read any letter that arrives in September carefully."]),

 ("faq",[
  ("How much will Medicare cost in 2027?",
   "The standard Part B premium is projected at about $209.50 per month for 2027, up from $202.90 in 2026, according to the 2026 Medicare Trustees Report. CMS confirms the actual figure in the fall. The standard Part D deductible rises to $700 from $615, and the annual Part D out-of-pocket cap rises to $2,400 from $2,100. What you personally pay also depends on your plan, your prescriptions and whether your income triggers the IRMAA surcharge."),
  ("What is the Part D out-of-pocket maximum for 2027?",
   "$2,400, up from $2,100 in 2026. Once your out-of-pocket spending on covered Part D drugs reaches that amount in a calendar year, you pay nothing more for covered drugs for the rest of the year. Most people never reach the cap; it exists to prevent a catastrophic year for people taking expensive medications."),
  ("Why is the Part D deductible going up so much?",
   "The standard Part D deductible rises from $615 to $700 for 2027, roughly 14 percent. The Part D parameters are adjusted annually and track growth in prescription drug costs across the program. For most people this is the more noticeable change of the two, because everyone filling covered prescriptions passes through the deductible, while relatively few reach the out-of-pocket cap."),
  ("Will my prescriptions be cheaper in 2027 because of drug price negotiation?",
   "Possibly, but not automatically. Negotiated prices for a second group of 15 Part D drugs take effect January 1, 2027, at an average of about 44 percent off list price. What you pay at the pharmacy is determined by your plan's formulary tier, your deductible and which phase of the benefit you are in, not directly by the list price. Check your specific drugs in your specific plan for 2027 during the Annual Enrollment Period."),
  ("When can I change my Medicare plan for 2027?",
   "The Annual Enrollment Period runs October 15 to December 7, 2026, with changes effective January 1, 2027. If you are in a Medicare Advantage plan, you also get one change between January 1 and March 31, 2027 during Medicare Advantage Open Enrollment. If your plan is being discontinued, you get a Special Enrollment Period and are not limited to those windows."),
 ]),
]),

# ============================================================ 1.3
A(
slug="medicare-drug-prices-2027-copay",
title="Medicare Drug Prices in 2027: Will Your Copay Drop?",
h1="Medicare's Negotiated Drug Prices Start January 1, 2027. Here Is Whether Your Copay Actually Changes.",
desc="A second round of 15 negotiated Part D drugs takes effect in 2027 at about 44% off list. Why that may not lower what you pay, and the four things that decide it.",
art="pricetag",
blocks=[
 ("p","You are going to see a lot of headlines this fall saying Medicare drug prices are dropping. They are not wrong. They are just answering a different question than the one you are asking."),
 ("p","The question you are asking is <strong>what you will pay at the pharmacy counter in January.</strong> That is not the same as the list price, and for a lot of people the answer will be \"about the same as now.\""),
 ("qa","On January 1, 2027, negotiated prices take effect for a second group of <strong>15 Part D drugs</strong>, averaging roughly <strong>44 percent off list price</strong>. But the negotiated price is what <strong>Medicare</strong> pays. What <strong>you</strong> pay is set by your plan's formulary tier, your deductible, and which phase of the benefit you are in. Some people will see a real drop. Others will see no change at all."),

 ("h2","What actually got negotiated"),
 ("p","Under the Inflation Reduction Act, Medicare negotiates prices directly with manufacturers for a growing list of high-spend drugs. The first 10 took effect in 2026. A second group of 15 follows on January 1, 2027, covering drugs used by about 5.3 million Part D enrollees."),
 ("p","Publicly confirmed drugs in this second group include the semaglutide medicines (Ozempic, Rybelsus and Wegovy), Trelegy Ellipta for COPD and asthma, Linzess, Ibrance and Xtandi for cancer, Austedo, Trulicity and Xifaxan. They treat diabetes, cardiovascular disease, obesity, respiratory disease, cancer and digestive conditions."),
 ("callout","One caution about lists you will see online.","Coverage of this program has been mixing up the negotiation rounds. The round taking effect in January 2027 is the second one. A third round was selected in January 2026 and does not take effect in 2027. If you see a drug list, check which year the prices actually apply to before you plan around it."),

 ("h2","The four things that decide what you actually pay"),
 ("h3","1. Which tier your plan puts the drug on"),
 ("p","This is the big one. Plans sort covered drugs into tiers, and your cost sharing follows the tier, not the list price. A drug can get a large negotiated discount and stay on a specialty tier where you owe a percentage of the cost. Tier placement is set by your plan and can change every January."),
 ("h3","2. Whether you have met your deductible"),
 ("p","The standard Part D deductible rises to $700 in 2027. Until you meet it, you generally pay the full negotiated price rather than a copay. A lower negotiated price genuinely helps here, because you are paying that price directly."),
 ("h3","3. Which phase of the benefit you are in"),
 ("p","Part D moves through phases across the year: the deductible, then initial coverage where you pay a copay or coinsurance, then the catastrophic phase once you hit the out-of-pocket cap, which is $2,400 in 2027. Where you sit in that sequence when you fill a prescription changes what you hand over."),
 ("h3","4. Whether your plan covers the drug at all"),
 ("p","A negotiated price applies to drugs a plan covers. Formularies change annually, and a plan can move, restrict or drop a drug for 2027 regardless of what was negotiated. Prior authorization and step therapy requirements can change too."),

 ("h2","Who is most likely to see a real drop"),
 ("ul",["People who spend part of the year in the deductible phase, where you pay the drug's price directly.",
        "People taking one of the negotiated drugs who do not reach the out-of-pocket cap, so their coinsurance is calculated on a lower price.",
        "People whose plan uses coinsurance, a percentage of cost, rather than a flat copay for that tier."]),
 ("p","Conversely, if you pay a flat $47 copay on a preferred brand tier, that copay is $47 whether the underlying price fell 44 percent or not. And if you already reach the out-of-pocket cap every year, you were going to stop paying at the cap either way."),

 ("h2","How to check your own answer during open enrollment"),
 ("ol",["Get your plan's 2027 formulary, which arrives with your Annual Notice of Change by September 30 or is posted on the plan's website.",
        "Look up each drug you take by name and note its tier for 2027, not 2026.",
        "Check whether the tier uses a flat copay or a percentage. Percentages move with price; flat copays do not.",
        "Check for new prior authorization or step therapy requirements on your drugs.",
        "Compare against at least one other plan available where you live, using the Medicare Plan Finder at Medicare.gov or a licensed agent who can run all of them at once."]),
 ("p","This is the comparison our licensed agents run every fall. For a closer look at how coverage rules play out for one much-asked-about drug class, our local Kentucky team wrote up <a href=\"https://www.bluegrassmedicarehelp.com/articles/does-medicare-cover-ozempic-wegovy-zepbound/\">how Medicare handles Ozempic, Wegovy and Zepbound</a>."),

 ("recap",["Negotiated prices for a second group of 15 Part D drugs take effect January 1, 2027, at roughly 44 percent off list price.",
           "The negotiated price is what Medicare pays. Your copay is set by your plan's tier, your deductible and your benefit phase.",
           "Flat copays do not move with list price. Coinsurance percentages do.",
           "Formularies change every January, so check the 2027 list rather than assuming last year's tiers.",
           "Published drug lists are mixing up the negotiation rounds; confirm which year a price actually applies to."]),

 ("faq",[
  ("Will Ozempic be cheaper on Medicare in 2027?",
   "The semaglutide medicines, sold as Ozempic, Rybelsus and Wegovy, are in the second group of drugs with negotiated Medicare prices effective January 1, 2027, at an average reduction across the group of about 44 percent off list price. Whether your own cost falls depends on your plan's formulary tier for 2027, whether you have met your deductible, and which phase of the Part D benefit you are in. Coverage also depends on what the drug is prescribed for. Check your specific plan's 2027 formulary."),
  ("Does a negotiated Medicare drug price lower my copay?",
   "Not automatically. The negotiated price sets what Medicare pays the manufacturer. Your out-of-pocket cost is determined by your plan's cost-sharing design: which tier the drug sits on, whether that tier charges a flat copay or a percentage of cost, whether you have met your deductible, and whether you have reached the annual out-of-pocket cap. If you pay a flat copay, a lower list price may not change your cost at all."),
  ("Which drugs have negotiated Medicare prices in 2027?",
   "A second group of 15 Part D drugs, used by roughly 5.3 million enrollees. Publicly confirmed among them are the semaglutide medicines (Ozempic, Rybelsus, Wegovy), Trelegy Ellipta, Linzess, Ibrance, Xtandi, Austedo, Trulicity and Xifaxan. They treat conditions including diabetes, cardiovascular disease, obesity, COPD and asthma, cancer and digestive disorders. Verify any list against CMS, because published coverage has been confusing this round with the later round selected in 2026."),
  ("What is the Part D out-of-pocket cap in 2027?",
   "$2,400, up from $2,100 in 2026. Once your out-of-pocket spending on covered Part D drugs reaches that amount in a calendar year, you pay nothing further for covered drugs that year. The standard Part D deductible for 2027 is $700, up from $615."),
  ("How do I find out what my drugs will cost next year?",
   "Your plan sends an Annual Notice of Change by September 30 with the coming year's formulary and cost sharing. Look up each drug by name, note the 2027 tier, and check whether that tier uses a flat copay or a percentage. Then compare against other plans available where you live using the Medicare Plan Finder at Medicare.gov during the Annual Enrollment Period, October 15 to December 7, or ask a licensed agent to run the comparison across every plan at once."),
 ]),
]),

# ============================================================ 1.4
A(
slug="stop-medicare-sales-calls",
title="How to Stop Medicare Sales Calls and Spot a Real Agent",
h1="Why Medicare Marketing Got More Aggressive This Year, and How to Tell an Advisor From a Lead Generator",
desc="Federal marketing rules loosened on October 1, 2026. What changed, what is still illegal, how to stop the calls, and how to tell a licensed agent from a lead seller.",
art="phone",
blocks=[
 ("p","\"Why do I get six of these calls a day, and how do they have my name?\" We get asked that more than almost anything else between October and December."),
 ("p","This year there is a specific answer, and it is not that you did something wrong. <strong>Several federal marketing rules changed on October 1, 2026</strong>, and most of them loosened."),
 ("qa","CMS rolled back several Medicare marketing guardrails for the 2027 plan year. The 48-hour waiting period before a sales appointment is gone, the required separation between educational and sales events is gone, and marketers may now use words like \"best\" if they can substantiate the claim. <strong>High-pressure sales tactics are still prohibited</strong>, and so is almost everything a scam call does. Knowing which is which is the useful skill."),

 ("h2","What actually changed on October 1"),
 ("ul",["<strong>The 48-hour rule is gone.</strong> An agent used to have to wait 48 hours after you signed a Scope of Appointment before discussing plans. Now that conversation can happen the same day.",
        "<strong>The separation between educational and sales events is gone.</strong> A 12-hour gap used to be required between an educational seminar and a marketing event at the same venue.",
        "<strong>Agents may collect a Scope of Appointment at an educational event.</strong> You can be asked at the seminar to sign up for a follow-up sales appointment.",
        "<strong>The disclaimer moved.</strong> The third-party marketing disclaimer must now be given before any discussion of plan benefits, rather than within the first 60 seconds of a call.",
        "<strong>Superlatives are allowed</strong> where the claim can be substantiated, so expect to see more \"best\" and \"top-rated\" language in advertising than last year."]),
 ("p","None of that makes an agent untrustworthy. It does mean the pace of a sales conversation can move faster than it used to, and that fewer built-in pauses exist between meeting someone and enrolling with them. The pause is yours to take now."),

 ("h2","What is still against the rules"),
 ("p","This list did not loosen, and it is the practical test for whether a call is legitimate:"),
 ("ul",["<strong>Nobody may cold-call you about Medicare Advantage or Part D</strong> without your permission. If you did not ask to be contacted, the call is not compliant.",
        "<strong>No door-to-door sales visits</strong> without an appointment you agreed to.",
        "<strong>No high-pressure tactics.</strong> That prohibition survived the rule changes intact.",
        "<strong>No enrolling you without your explicit consent.</strong>",
        "<strong>Medicare will not call you</strong> to ask for your Medicare number, your Social Security number or your bank details. Neither will a legitimate agent, out of the blue."]),
 ("callout","The single most useful rule to remember.","If you did not initiate the contact, do not give out your Medicare number. Hang up and call the person or company back on a number you looked up yourself. A real agent will not mind at all."),

 ("h2","Why they have your name"),
 ("p","Most of those calls do not come from insurance agents. They come from lead generators: companies whose product is your contact information, which they sell to multiple agencies at once. That is why answering one call produces six more."),
 ("h3","Where the lists come from"),
 ("p","The usual sources are a \"free Medicare guide\" or benefits-checker form you filled out online, a mailer with a reply card, a sweepstakes entry, or a data broker who bought a list that includes your age. Anything that asks for your phone number in exchange for a Medicare quote is, functionally, a lead form."),
 ("h3","Why answering makes it worse"),
 ("p","Lead lists are scored. A number that answers, stays on the line, or presses a key gets marked as live and resold at a higher price. The single quietest thing you can do with an unknown Medicare call is not answer it at all, and let voicemail sort the real callers from the rest."),

 ("h2","Five things that actually reduce the calls"),
 ("ol",["<strong>Register on the National Do Not Call Registry</strong> at donotcall.gov. It will not stop scam callers who already ignore the law, but it removes you from compliant marketing lists.",
        "<strong>Say the words \"put me on your do-not-call list\" and note the date.</strong> Legitimate companies must honor that request. It gives you standing if they call again.",
        "<strong>Stop filling in online Medicare quote forms.</strong> One submission can be sold to a dozen buyers. If you want a comparison, call an agency directly instead.",
        "<strong>Do not press any key, including the one that promises to remove you.</strong> On an illegal robocall, pressing a key confirms a live human answered and increases the calls.",
        "<strong>Report the ones that break the rules.</strong> Call 1-800-MEDICARE for Medicare-specific complaints, or file with the FTC at reportfraud.ftc.gov."]),

 ("h2","Advisor or lead generator: how to tell in one call"),
 ("table","What to listen for",
   ["A licensed agent","A lead generator or scam call"],
   [["Gives you a name, an agency and a license or National Producer Number when asked","Deflects, or only names a vague \"Medicare benefits center\""],
    ["Can tell you which carriers they represent, and admits they do not represent all of them","Claims to offer everything, or claims to be from Medicare"],
    ["Asks about your doctors, your prescriptions and your budget before mentioning a plan","Asks for your Medicare number early"],
    ["Is willing to be called back on a number you look up","Pressures you to decide on this call"],
    ["Says \"keep what you have\" when that is the right answer","Always finds a reason to switch you"]]),
 ("p","One more that costs nothing: ask whether they will still be your point of contact next year. A lead generator cannot answer that question, because they were never going to be."),
 ("p","For a fuller look at the scams that specifically target Medicare beneficiaries and how to report them, our local Kentucky team put together a guide to <a href=\"https://www.bluegrassmedicarehelp.com/articles/medicare-scams-how-to-protect-yourself/\">protecting your Medicare number</a>."),

 ("recap",["Federal marketing rules loosened on October 1, 2026, so this enrollment season moves faster than last year's.",
           "Cold calls, door-to-door sales and high-pressure tactics are all still prohibited.",
           "Most calls come from lead generators selling your contact information, not from agents.",
           "Register at donotcall.gov, ask to be added to internal do-not-call lists, and stop filling in online quote forms.",
           "Never give your Medicare number to someone who called you first."]),

 ("faq",[
  ("Why am I getting so many Medicare phone calls?",
   "Most of them come from lead generation companies rather than insurance agents. Their product is your contact information, which they sell to several agencies at once, so a single online form or reply card can produce calls from many different callers. Volume also rises sharply between October 15 and December 7, the Annual Enrollment Period, and federal marketing rules loosened on October 1, 2026, which allows a faster sales process than in previous years."),
  ("Is it legal for someone to cold-call me about Medicare?",
   "No. Federal rules prohibit unsolicited contact about Medicare Advantage and Part D plans, including cold calls and unrequested door-to-door visits, unless you gave permission to be contacted. Filling out an online form or returning a reply card often counts as giving that permission. High-pressure sales tactics remain prohibited even after the October 2026 rule changes."),
  ("How do I stop Medicare marketing calls?",
   "Register your number at donotcall.gov, tell callers directly to add you to their internal do-not-call list and note the date, and stop submitting online Medicare quote forms, which are typically sold to multiple buyers. Do not press any key on a robocall, including one offering to remove you, because that confirms a live person answered. Report violations to 1-800-MEDICARE or the FTC at reportfraud.ftc.gov."),
  ("How can I tell if a Medicare agent is legitimate?",
   "Ask for their name, their agency and their license or National Producer Number, and ask which carriers they represent. A licensed agent will answer all of it, will acknowledge that they do not represent every company, and will let you call them back on a number you look up yourself. They will also ask about your doctors, prescriptions and budget before recommending anything. Someone who asks for your Medicare number early, claims to be calling from Medicare, or pressures you to decide on the call is neither licensed nor safe."),
  ("Will Medicare ever call me?",
   "Medicare does not call beneficiaries to ask for a Medicare number, Social Security number or bank details, and does not sell plans over the phone. If someone claims to be from Medicare and asks for that information, hang up and call 1-800-MEDICARE directly using the number on your card or on Medicare.gov."),
 ]),
]),

# ============================================================ 1.5
A(
slug="medicare-part-d-2027",
title="Medicare Part D in 2027: Deductible, Cap, and Costs",
h1="What Your Medicare Drug Plan Costs in 2027: the $700 Deductible and the $2,400 Cap",
desc="How Part D works in 2027: the $700 deductible, the $2,400 out-of-pocket cap, the phases in between, and the payment plan that spreads the cost.",
art="stairs",
blocks=[
 ("p","A woman called our office in February a couple of years ago, upset. Her drug plan had not changed, her prescriptions had not changed, and her January pharmacy bill was four times what she paid in December."),
 ("p","Nothing had gone wrong. She had simply started a new plan year, and the deductible reset. It is the single most common January phone call in this business, and it is entirely predictable once you know how the benefit is built."),
 ("qa","Part D moves through <strong>phases</strong> across a calendar year and your cost changes as you move. For 2027 the standard deductible is <strong>$700</strong> and the annual out-of-pocket cap is <strong>$2,400</strong>. You pay the most in January and least in December, and once you reach the cap you pay nothing more for covered drugs that year."),

 ("h2","The phases, in order"),
 ("h3","The deductible phase"),
 ("p","At the start of the year you pay the full negotiated cost of your drugs until you have spent the deductible. The 2027 standard is <strong>$700</strong>, up from $615. Plans can set a lower deductible or none at all, and many exclude generics from it entirely, so read your own plan rather than assuming the standard applies to you."),
 ("h3","The initial coverage phase"),
 ("p","After the deductible, you pay a copay (a flat dollar amount) or coinsurance (a percentage) and the plan pays the rest. This is where most people spend most of the year."),
 ("h3","The catastrophic phase"),
 ("p","Once your out-of-pocket spending reaches <strong>$2,400</strong>, you are done paying for covered drugs for the rest of the calendar year. This cap has only existed in this form since 2025 and it replaced a much harsher structure. If you take an expensive drug, it is the most valuable feature of the entire benefit."),
 ("callout","What counts toward the cap, and what does not.","Your deductible and your copays count. Your monthly premium does not. Drugs your plan does not cover do not count either, which is why formulary checks matter more than premium comparisons."),

 ("h2","Why January costs the most"),
 ("p","Every phase resets on January 1. Someone who finished December in the catastrophic phase paying nothing starts January paying full price until the new deductible is met. Nothing about the plan changed. The calendar did."),
 ("p","This is worth planning for rather than being surprised by, particularly if you take a drug that costs several hundred dollars a month. Which brings us to the option most people have never heard of."),

 ("h2","The payment plan almost nobody uses"),
 ("p","The <strong>Medicare Prescription Payment Plan</strong> lets you spread your out-of-pocket drug costs across the remaining months of the year in level payments, rather than paying at the pharmacy counter. It is free to join, available with any Part D plan, and you opt in."),
 ("p","Two things to be honest about. It does <strong>not</strong> reduce what you owe in total; it changes the timing. And if you join late in the year, the same balance gets divided across fewer months, so the payments are larger."),
 ("p","It helps most for someone facing a large bill early in the year who would rather pay it in level installments than all at once in January. It helps least for someone with modest, steady drug costs, who is generally better off just paying at the counter."),

 ("h2","What decides your own number"),
 ("ul",["<strong>Which tier your plan puts each drug on.</strong> Tiers drive cost sharing far more than the plan's premium does.",
        "<strong>Copay or coinsurance.</strong> A flat copay is predictable. A percentage moves with the drug's price.",
        "<strong>Whether your pharmacy is preferred.</strong> Most plans have preferred and standard pharmacies, and the difference on the same drug can be substantial.",
        "<strong>Restrictions.</strong> Prior authorization, step therapy and quantity limits can delay or block a fill even when the drug is technically covered.",
        "<strong>Whether you qualify for Extra Help.</strong> This is the low-income subsidy, and it changes the math completely. It is worth checking even if you assume you earn too much."]),

 ("h2","What to do during open enrollment"),
 ("ol",["List every prescription you take, with the dose.",
        "Look each one up in your plan's 2027 formulary and note the tier and whether it carries restrictions.",
        "Check whether your regular pharmacy is preferred under the plan for 2027.",
        "Compare total annual cost, meaning premium plus deductible plus expected copays, not premium alone. The cheapest premium is frequently not the cheapest year.",
        "If January is going to be expensive, decide before it arrives whether the payment plan is worth using."]),
 ("p","Our licensed agents run drug-by-drug comparisons across every plan available in your area at no cost. For a plain-English walkthrough of how the parts fit together, our local Kentucky team wrote up <a href=\"https://www.bluegrassmedicarehelp.com/articles/medicare-part-d-prescription-drug-plans/\">how Part D drug plans work</a>."),

 ("recap",["The 2027 standard Part D deductible is $700, up from $615, and the out-of-pocket cap is $2,400, up from $2,100.",
           "Part D runs in phases and every phase resets January 1, which is why January costs the most.",
           "Premiums do not count toward the out-of-pocket cap. Deductibles and copays do.",
           "The Medicare Prescription Payment Plan spreads costs across the year but does not reduce them.",
           "Compare total annual cost, not premium, and check tiers, pharmacy status and restrictions."]),

 ("faq",[
  ("What is the Medicare Part D deductible for 2027?",
   "The standard deductible is $700, up from $615 in 2026. That is the maximum a plan may charge; plans can set a lower deductible or none at all, and many apply no deductible to generic drugs. Check your own plan's Annual Notice of Change or Summary of Benefits rather than assuming the standard figure applies to you."),
  ("What is the Part D out-of-pocket maximum in 2027?",
   "$2,400, up from $2,100 in 2026. Once your out-of-pocket spending on covered drugs reaches that amount in a calendar year, you pay nothing more for covered Part D drugs for the rest of that year. Your deductible and copays count toward the cap. Your monthly premium does not, and neither does spending on drugs your plan does not cover."),
  ("Why did my prescription cost so much more in January?",
   "Because the Part D benefit resets on January 1. If you finished the previous year in the catastrophic phase paying nothing for covered drugs, you begin January in the deductible phase paying full negotiated cost until you have spent the deductible, which is $700 in 2027. Nothing about your plan has to change for the January bill to be several times the December bill."),
  ("What is the Medicare Prescription Payment Plan?",
   "It is a free option that lets you spread your out-of-pocket Part D drug costs across the remaining months of the calendar year in level monthly payments instead of paying the full amount at the pharmacy. It is available with any Part D plan and you have to opt in. It does not lower your total cost, only the timing, and joining later in the year means the same balance is divided across fewer months, so the payments are larger."),
  ("Does the cheapest Part D premium mean the lowest cost?",
   "Often not. A plan's premium is one of several inputs. The deductible, which tier each of your drugs falls on, whether the tier charges a flat copay or a percentage, whether your pharmacy is preferred, and any prior authorization or step therapy requirements all affect what you actually spend. Compare total expected annual cost for the specific drugs you take rather than comparing premiums."),
 ]),
]),
]


# ---------------------------------------------------------------- illustrations
def svg(kind):
    """1200x630 TIG-palette illustration. Concept over product, compliance-safe."""
    head = (f'<svg xmlns="http://www.w3.org/2000/svg" width="1200" height="630" '
            f'viewBox="0 0 1200 630" role="img" aria-label="Tyler Insurance Group illustration">'
            f'<rect width="1200" height="630" fill="{GOLD_T}"/>')
    tail = '</svg>'
    F = ('font-family="Source Sans 3, Helvetica, Arial, sans-serif"')

    if kind == "letter":                      # envelope + two doors + 63-day badge
        b = f'''
<circle cx="1050" cy="110" r="150" fill="{GOLD}" opacity=".30"/>
<circle cx="140" cy="560" r="120" fill="{GRAY}" opacity=".16"/>
<g transform="translate(90,175)">
  <rect x="0" y="40" width="360" height="240" rx="14" fill="{WHITE}" stroke="{GRAY_D}" stroke-width="4"/>
  <path d="M0 54 L180 190 L360 54" fill="none" stroke="{GRAY_D}" stroke-width="4"/>
  <rect x="52" y="0" width="256" height="150" rx="10" fill="{WHITE}" stroke="{GOLD_D}" stroke-width="4"/>
  <rect x="84" y="34" width="150" height="12" rx="6" fill="{GRAY}"/>
  <rect x="84" y="62" width="192" height="10" rx="5" fill="{GOLD}"/>
  <rect x="84" y="86" width="120" height="10" rx="5" fill="{GOLD}"/>
</g>
<g transform="translate(560,150)">
  <rect x="0" y="60" width="230" height="290" rx="12" fill="{GOLD}" stroke="{GOLD_D}" stroke-width="4"/>
  <circle cx="115" cy="300" r="11" fill="{WHITE}"/>
  <text x="115" y="150" {F} font-size="30" font-weight="700" fill="{WHITE}" text-anchor="middle">Medigap</text>
  <text x="115" y="188" {F} font-size="21" fill="{WHITE}" text-anchor="middle">no health</text>
  <text x="115" y="214" {F} font-size="21" fill="{WHITE}" text-anchor="middle">questions</text>
</g>
<g transform="translate(840,150)">
  <rect x="0" y="60" width="230" height="290" rx="12" fill="{GRAY}" stroke="{GRAY_D}" stroke-width="4"/>
  <circle cx="115" cy="300" r="11" fill="{WHITE}"/>
  <text x="115" y="160" {F} font-size="28" font-weight="700" fill="{WHITE}" text-anchor="middle">Another</text>
  <text x="115" y="196" {F} font-size="28" font-weight="700" fill="{WHITE}" text-anchor="middle">Advantage</text>
  <text x="115" y="230" {F} font-size="28" font-weight="700" fill="{WHITE}" text-anchor="middle">plan</text>
</g>
<g transform="translate(600,528)">
  <rect x="0" y="0" width="330" height="70" rx="35" fill="{WHITE}" stroke="{GOLD_D}" stroke-width="5"/>
  <text x="165" y="47" {F} font-size="32" font-weight="700" fill="{GRAY_D}" text-anchor="middle">63 days to act</text>
</g>'''
        return head + b + tail

    if kind == "calendar":                    # 2027 calendar + three change chips
        b = f'''
<circle cx="120" cy="90" r="120" fill="{GOLD}" opacity=".28"/>
<circle cx="1090" cy="560" r="140" fill="{GRAY}" opacity=".14"/>
<g transform="translate(95,140)">
  <rect x="0" y="0" width="330" height="350" rx="18" fill="{WHITE}" stroke="{GRAY_D}" stroke-width="4"/>
  <rect x="0" y="0" width="330" height="86" rx="18" fill="{GOLD}"/>
  <rect x="0" y="68" width="330" height="18" fill="{GOLD}"/>
  <text x="165" y="62" {F} font-size="46" font-weight="700" fill="{WHITE}" text-anchor="middle">2027</text>
  <rect x="70" y="-18" width="22" height="46" rx="11" fill="{GRAY_D}"/>
  <rect x="238" y="-18" width="22" height="46" rx="11" fill="{GRAY_D}"/>
  <g fill="{GOLD_T}">
    <rect x="34" y="120" width="52" height="42" rx="8"/><rect x="104" y="120" width="52" height="42" rx="8"/>
    <rect x="174" y="120" width="52" height="42" rx="8"/><rect x="244" y="120" width="52" height="42" rx="8"/>
    <rect x="34" y="180" width="52" height="42" rx="8"/><rect x="174" y="180" width="52" height="42" rx="8"/>
    <rect x="244" y="180" width="52" height="42" rx="8"/>
    <rect x="34" y="240" width="52" height="42" rx="8"/><rect x="104" y="240" width="52" height="42" rx="8"/>
    <rect x="174" y="240" width="52" height="42" rx="8"/><rect x="244" y="240" width="52" height="42" rx="8"/>
  </g>
  <rect x="104" y="180" width="52" height="42" rx="8" fill="{GOLD}" stroke="{GOLD_D}" stroke-width="3"/>
  <text x="130" y="209" {F} font-size="24" font-weight="700" fill="{WHITE}" text-anchor="middle">1</text>
</g>
<g transform="translate(510,140)">
  <rect x="0" y="0" width="590" height="94" rx="16" fill="{WHITE}" stroke="{GOLD_D}" stroke-width="4"/>
  <text x="34" y="42" {F} font-size="25" fill="{GRAY_D}">Part B premium</text>
  <text x="34" y="76" {F} font-size="31" font-weight="700" fill="{GRAY_D}">about $209.50 projected</text>
  <rect x="0" y="128" width="590" height="94" rx="16" fill="{WHITE}" stroke="{GOLD_D}" stroke-width="4"/>
  <text x="34" y="170" {F} font-size="25" fill="{GRAY_D}">Part D deductible</text>
  <text x="34" y="204" {F} font-size="31" font-weight="700" fill="{GRAY_D}">$615 rises to $700</text>
  <rect x="0" y="256" width="590" height="94" rx="16" fill="{GOLD}" stroke="{GOLD_D}" stroke-width="4"/>
  <text x="34" y="298" {F} font-size="25" fill="{WHITE}">Part D out-of-pocket cap</text>
  <text x="34" y="332" {F} font-size="31" font-weight="700" fill="{WHITE}">$2,100 rises to $2,400</text>
</g>
<text x="600" y="560" {F} font-size="30" fill="{GRAY_D}" text-anchor="middle">What changes on January 1</text>'''
        return head + b + tail

    if kind == "pricetag":                    # list price down, copay question mark
        b = f'''
<circle cx="1080" cy="120" r="140" fill="{GOLD}" opacity=".28"/>
<circle cx="120" cy="540" r="110" fill="{GRAY}" opacity=".15"/>
<g transform="translate(120,180)">
  <path d="M20 0 H210 A20 20 0 0 1 230 20 V150 A20 20 0 0 1 210 170 H20 L-70 85 Z"
        fill="{GOLD}" stroke="{GOLD_D}" stroke-width="4"/>
  <circle cx="30" cy="85" r="16" fill="{WHITE}"/>
  <text x="130" y="76" {F} font-size="23" fill="{WHITE}" text-anchor="middle">list price</text>
  <text x="130" y="124" {F} font-size="40" font-weight="700" fill="{WHITE}" text-anchor="middle">-44%</text>
</g>
<g transform="translate(470,268)">
  <path d="M0 22 H140" stroke="{GRAY_D}" stroke-width="7" stroke-linecap="round"/>
  <path d="M112 0 L146 22 L112 44" fill="none" stroke="{GRAY_D}" stroke-width="7"
        stroke-linecap="round" stroke-linejoin="round"/>
</g>
<g transform="translate(680,150)">
  <rect x="0" y="0" width="330" height="330" rx="18" fill="{WHITE}" stroke="{GRAY_D}" stroke-width="4"/>
  <rect x="0" y="0" width="330" height="72" rx="18" fill="{GRAY}"/>
  <rect x="0" y="54" width="330" height="18" fill="{GRAY}"/>
  <text x="165" y="50" {F} font-size="27" font-weight="700" fill="{WHITE}" text-anchor="middle">your copay</text>
  <text x="165" y="230" {F} font-size="140" font-weight="700" fill="{GOLD_D}" text-anchor="middle">?</text>
  <text x="165" y="278" {F} font-size="21" fill="{GRAY_D}" text-anchor="middle">set by your plan's tier,</text>
  <text x="165" y="304" {F} font-size="21" fill="{GRAY_D}" text-anchor="middle">deductible and phase</text>
</g>
<text x="600" y="560" {F} font-size="30" fill="{GRAY_D}" text-anchor="middle">A lower negotiated price is not automatically a lower copay</text>'''
        return head + b + tail

    if kind == "phone":                       # ringing phone, shield, credential badge
        b = f'''
<circle cx="1070" cy="120" r="140" fill="{GOLD}" opacity=".28"/>
<circle cx="130" cy="560" r="120" fill="{GRAY}" opacity=".15"/>
<g transform="translate(140,170)">
  <rect x="0" y="0" width="190" height="300" rx="26" fill="{GRAY}" stroke="{GRAY_D}" stroke-width="4"/>
  <rect x="16" y="34" width="158" height="230" rx="10" fill="{WHITE}"/>
  <circle cx="95" cy="282" r="10" fill="{WHITE}" opacity=".8"/>
  <text x="95" y="130" {F} font-size="22" fill="{GRAY_D}" text-anchor="middle">Unknown</text>
  <text x="95" y="162" {F} font-size="22" fill="{GRAY_D}" text-anchor="middle">Caller</text>
  <g stroke="{GOLD_D}" stroke-width="7" fill="none" stroke-linecap="round">
    <path d="M-42 96 A 62 62 0 0 0 -42 204"/>
    <path d="M-74 74 A 104 104 0 0 0 -74 226"/>
    <path d="M232 96 A 62 62 0 0 1 232 204"/>
    <path d="M264 74 A 104 104 0 0 1 264 226"/>
  </g>
</g>
<g transform="translate(640,150)">
  <path d="M150 0 L292 52 V190 C292 268 226 312 150 336 C74 312 8 268 8 190 V52 Z"
        fill="{GOLD}" stroke="{GOLD_D}" stroke-width="4"/>
  <rect x="72" y="120" width="156" height="46" rx="10" fill="{WHITE}"/>
  <rect x="72" y="182" width="156" height="16" rx="8" fill="{WHITE}" opacity=".85"/>
  <rect x="72" y="212" width="110" height="16" rx="8" fill="{WHITE}" opacity=".85"/>
  <text x="150" y="153" {F} font-size="24" font-weight="700" fill="{GRAY_D}" text-anchor="middle">LICENSED</text>
  <circle cx="150" cy="272" r="26" fill="{WHITE}"/>
  <path d="M138 272 l9 9 17-18" fill="none" stroke="{GOLD_D}" stroke-width="6"
        stroke-linecap="round" stroke-linejoin="round"/>
</g>
<text x="600" y="560" {F} font-size="30" fill="{GRAY_D}" text-anchor="middle">Ask for a name, an agency and a license number</text>'''
        return head + b + tail

    if kind == "stairs":                      # four phases with the two key numbers
        b = f'''
<circle cx="1080" cy="110" r="130" fill="{GOLD}" opacity=".28"/>
<circle cx="110" cy="560" r="110" fill="{GRAY}" opacity=".14"/>
<g transform="translate(120,150)">
  <rect x="0" y="250" width="220" height="130" rx="12" fill="{GRAY}" stroke="{GRAY_D}" stroke-width="4"/>
  <text x="110" y="300" {F} font-size="23" fill="{WHITE}" text-anchor="middle">Deductible</text>
  <text x="110" y="342" {F} font-size="38" font-weight="700" fill="{WHITE}" text-anchor="middle">$700</text>

  <rect x="248" y="170" width="220" height="210" rx="12" fill="{GOLD}" stroke="{GOLD_D}" stroke-width="4"/>
  <text x="358" y="228" {F} font-size="23" fill="{WHITE}" text-anchor="middle">Initial</text>
  <text x="358" y="258" {F} font-size="23" fill="{WHITE}" text-anchor="middle">coverage</text>
  <text x="358" y="308" {F} font-size="26" font-weight="700" fill="{WHITE}" text-anchor="middle">copay or</text>
  <text x="358" y="338" {F} font-size="26" font-weight="700" fill="{WHITE}" text-anchor="middle">coinsurance</text>

  <rect x="496" y="80" width="220" height="300" rx="12" fill="{WHITE}" stroke="{GOLD_D}" stroke-width="4"/>
  <text x="606" y="140" {F} font-size="23" fill="{GRAY_D}" text-anchor="middle">Cap reached</text>
  <text x="606" y="196" {F} font-size="44" font-weight="700" fill="{GOLD_D}" text-anchor="middle">$2,400</text>
  <text x="606" y="252" {F} font-size="23" fill="{GRAY_D}" text-anchor="middle">then you pay</text>
  <text x="606" y="300" {F} font-size="40" font-weight="700" fill="{GRAY_D}" text-anchor="middle">$0</text>
</g>
<g transform="translate(120,545)">
  <path d="M0 12 H830" stroke="{GRAY_D}" stroke-width="5" stroke-linecap="round"/>
  <path d="M806 0 L836 12 L806 24" fill="none" stroke="{GRAY_D}" stroke-width="5"
        stroke-linecap="round" stroke-linejoin="round"/>
  <text x="12" y="52" {F} font-size="24" fill="{GRAY_D}">January</text>
  <text x="836" y="52" {F} font-size="24" fill="{GRAY_D}" text-anchor="end">December</text>
</g>
<text x="985" y="330" {F} font-size="26" fill="{GRAY_D}" text-anchor="middle">Every phase</text>
<text x="985" y="364" {F} font-size="26" fill="{GRAY_D}" text-anchor="middle">resets Jan 1</text>'''
        return head + b + tail

    raise ValueError(kind)


# ---------------------------------------------------------------- HTML
def esc(s):
    return _html.escape(s, quote=False)

def html_doc(a):
    P = []
    for b in a["blocks"]:
        k = b[0]
        if k == "p":     P.append(f"<p>{b[1]}</p>")
        elif k == "h2":  P.append(f"<h2>{esc(b[1])}</h2>")
        elif k == "h3":  P.append(f"<h3>{esc(b[1])}</h3>")
        elif k == "ul":  P.append("<ul>" + "".join(f"<li>{i}</li>" for i in b[1]) + "</ul>")
        elif k == "ol":  P.append("<ol>" + "".join(f"<li>{i}</li>" for i in b[1]) + "</ol>")
        elif k == "qa":
            P.append(f'<div class="qa"><p><strong>Quick answer.</strong> {b[1]}</p></div>')
        elif k == "callout":
            P.append(f'<div class="callout"><p><strong>{esc(b[1])}</strong> {b[2]}</p></div>')
        elif k == "table":
            cap, hdr, rows = b[1], b[2], b[3]
            th = "".join(f"<th>{esc(h)}</th>" for h in hdr)
            tr = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in r) + "</tr>" for r in rows)
            P.append(f'<table><caption>{esc(cap)}</caption><thead><tr>{th}</tr></thead><tbody>{tr}</tbody></table>')
        elif k == "recap":
            P.append("<h2>Quick recap</h2><ul>" + "".join(f"<li>{i}</li>" for i in b[1]) + "</ul>")
        elif k == "faq":
            P.append("<h2>Frequently asked questions</h2>")
            for q, ans in b[1]:
                P.append(f"<h3>{esc(q)}</h3><p>{ans}</p>")
    body = "\n".join(P)
    return f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(a['title'])}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,600&family=Source+Sans+3:wght@400;600;700&display=swap" rel="stylesheet">
<style>
:root{{--gold:{GOLD};--goldd:{GOLD_D};--goldt:{GOLD_T};--gray:{GRAY};--grayd:{GRAY_D};--ink:{INK};--mute:{MUTE};--line:{LINE};}}
*{{box-sizing:border-box}}
body{{margin:0;background:#fff;color:var(--ink);font-family:'Source Sans 3',Helvetica,Arial,sans-serif;font-size:18px;line-height:1.85}}
.note{{background:var(--goldt);border:1px solid var(--goldd);border-radius:10px;padding:14px 18px;margin:24px auto;max-width:820px;font-size:14.5px;color:var(--grayd);line-height:1.6}}
.note b{{font-family:Fraunces,Georgia,serif}}
main{{max-width:820px;margin:0 auto;padding:8px 28px 64px}}
h1{{font-family:Fraunces,Georgia,serif;font-weight:600;font-size:38px;line-height:1.18;margin:28px 0 10px}}
h2{{font-family:Fraunces,Georgia,serif;font-weight:600;font-size:27px;line-height:1.25;margin:44px 0 12px}}
h3{{font-family:Fraunces,Georgia,serif;font-weight:600;font-size:21px;margin:32px 0 8px}}
p{{margin:0 0 26px}}
.byline{{color:var(--mute);font-size:15.5px;margin:0 0 30px;padding-bottom:18px;border-bottom:2px solid var(--gold)}}
ul,ol{{margin:0 0 26px;padding-left:24px}} li{{margin-bottom:12px}}
.qa{{background:var(--gold);border-radius:12px;padding:22px 26px;margin:0 0 30px}}
.qa p{{margin:0;color:#fff;font-size:18.5px}} .qa strong{{color:#fff}}
.callout{{background:var(--goldt);border-left:5px solid var(--goldd);border-radius:0 12px 12px 0;padding:20px 24px;margin:0 0 30px}}
.callout p{{margin:0}}
table{{border-collapse:collapse;width:100%;margin:0 0 30px;font-size:16.5px}}
caption{{caption-side:top;text-align:left;font-size:14.5px;color:var(--mute);padding-bottom:8px}}
th{{background:var(--gray);color:#fff;text-align:left;padding:12px 14px;font-weight:600}}
td{{border-bottom:1px solid var(--line);padding:12px 14px;vertical-align:top;line-height:1.7}}
a{{color:var(--goldd);font-weight:600}}
.cta{{background:var(--goldt);border:1px solid var(--goldd);border-radius:12px;padding:24px 26px;margin:40px 0 0}}
.cta h2{{margin-top:0}}
.author{{border-top:1px solid var(--line);margin-top:40px;padding-top:18px;font-style:italic;color:var(--mute);font-size:15.5px}}
.disc{{color:var(--mute);font-size:14px;line-height:1.7;margin-top:24px}}
</style></head><body>
<div class="note"><b>How to use this file.</b><br>
Copy the title into the TIG editor's Title field. Then select everything from the byline down and
paste it into Blog Content. Upload the matching illustration separately via Add Media. Headings in
the body are H2 and H3 only, because the Title field becomes the page H1.</div>
<main>
<h1>{esc(a['h1'])}</h1>
<p class="byline">By Austin Tyler &middot; Tyler Insurance Group &middot; Updated {UPDATED}</p>
{body}
<div class="cta"><h2>Talk it through with a licensed agent</h2>
<p>Our licensed agents compare every plan available where you live, check your doctors and your
prescriptions by name, and tell you plainly when the plan you already have is the right one. There
is no cost to work with us.</p>
<p><strong>Call Tyler Insurance Group: {PHONE}</strong></p></div>
<p class="author">{esc(AUTHOR_LINE)}</p>
<p class="disc">{esc(DISCLAIMER)}</p>
</main></body></html>"""


# ---------------------------------------------------------------- DOCX
def docx_doc(a, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    pf = st.paragraph_format
    pf.line_spacing = 1.5; pf.space_after = Pt(15)

    def shade(p, hexcolor):
        el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexcolor.lstrip("#")); p._p.get_or_add_pPr().append(el)

    INLINE = re.compile(r'(<strong>.*?</strong>|<em>.*?</em>|<a [^>]*>.*?</a>)', re.S)
    def runs(p, text, color=None):
        for part in INLINE.split(text):
            if not part: continue
            if part.startswith("<strong>"):
                r = p.add_run(re.sub(r"</?strong>", "", part)); r.bold = True
            elif part.startswith("<em>"):
                r = p.add_run(re.sub(r"</?em>", "", part)); r.italic = True
            elif part.startswith("<a "):
                url = re.search(r'href="([^"]+)"', part).group(1)
                txt = re.sub(r"<[^>]+>", "", part)
                r = p.add_run(f"{txt} ({url})"); r.underline = True
            else:
                r = p.add_run(_html.unescape(re.sub(r"<[^>]+>", "", part)))
            if color: r.font.color.rgb = color
        return p

    # instruction line, plain-bold title, second instruction line
    g = doc.add_paragraph(); r = g.add_run("Copy the line below into the Title field.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
    t = doc.add_paragraph(); r = t.add_run(a["title"]); r.bold = True; r.font.size = Pt(14)
    g2 = doc.add_paragraph(); r = g2.add_run("Body starts below. Select from the byline down and paste into Blog Content.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)

    by = doc.add_paragraph()
    r = by.add_run(f"By Austin Tyler · Tyler Insurance Group · Updated {UPDATED}")
    r.italic = True; r.font.color.rgb = RGBColor(0x5F, 0x59, 0x4F)

    for b in a["blocks"]:
        k = b[0]
        if k == "p":     runs(doc.add_paragraph(), b[1])
        elif k == "h2":  doc.add_heading(b[1], level=2)
        elif k == "h3":  doc.add_heading(b[1], level=3)
        elif k == "ul":
            for i in b[1]: runs(doc.add_paragraph(style="List Bullet"), i)
        elif k == "ol":
            for i in b[1]: runs(doc.add_paragraph(style="List Number"), i)
        elif k == "qa":
            p = doc.add_paragraph(); shade(p, GOLD_T)
            p.add_run("Quick answer. ").bold = True; runs(p, b[1])
        elif k == "callout":
            p = doc.add_paragraph(); shade(p, GOLD_T)
            p.add_run(b[1] + " ").bold = True; runs(p, b[2])
        elif k == "table":
            cap, hdr, rows = b[1], b[2], b[3]
            c = doc.add_paragraph(); r = c.add_run(cap)
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
            tb = doc.add_table(rows=1, cols=len(hdr)); tb.style = "Table Grid"
            for i, h in enumerate(hdr):
                cell = tb.rows[0].cells[i]; cell.text = ""
                rr = cell.paragraphs[0].add_run(h); rr.bold = True
            for row in rows:
                cells = tb.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = ""; runs(cells[i].paragraphs[0], v)
        elif k == "recap":
            doc.add_heading("Quick recap", level=2)
            for i in b[1]: runs(doc.add_paragraph(style="List Bullet"), i)
        elif k == "faq":
            doc.add_heading("Frequently asked questions", level=2)
            for q, ans in b[1]:
                doc.add_heading(q, level=3); runs(doc.add_paragraph(), ans)

    doc.add_heading("Talk it through with a licensed agent", level=2)
    runs(doc.add_paragraph(),
         "Our licensed agents compare every plan available where you live, check your doctors and "
         "your prescriptions by name, and tell you plainly when the plan you already have is the "
         "right one. There is no cost to work with us.")
    p = doc.add_paragraph(); p.add_run(f"Call Tyler Insurance Group: {PHONE}").bold = True

    p = doc.add_paragraph(); r = p.add_run(AUTHOR_LINE)
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x5F, 0x59, 0x4F)
    pb = OxmlElement("w:pBdr"); tp = OxmlElement("w:top")
    tp.set(qn("w:val"), "single"); tp.set(qn("w:sz"), "6"); tp.set(qn("w:color"), "CCCCCC")
    pb.append(tp); p._p.get_or_add_pPr().append(pb)

    p = doc.add_paragraph(); r = p.add_run(DISCLAIMER)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    doc.save(path)


# ---------------------------------------------------------------- build
def main():
    os.makedirs(OUT, exist_ok=True)
    import xml.dom.minidom, cairosvg
    made = []
    for a in ARTICLES:
        s = a["slug"]
        open(f"{OUT}/{s}.html", "w", encoding="utf-8").write(html_doc(a))
        docx_doc(a, f"{OUT}/{s}.docx")
        sv = svg(a["art"])
        xml.dom.minidom.parseString(sv)                       # well-formedness
        open(f"{OUT}/{s}.svg", "w", encoding="utf-8").write(sv)
        cairosvg.svg2png(bytestring=sv.encode(), write_to=f"{OUT}/{s}.png",
                         output_width=1200, output_height=630)
        made.append(s)
        print(f"  built {s}")
    print(f"\n{len(made)} articles -> {OUT}/")

if __name__ == "__main__":
    main()
